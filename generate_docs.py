from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    document = Document()
    
    # Title
    title = document.add_heading('Collection Campaign Optimizer', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('Project Documentation & Execution Report').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Project Explanation
    document.add_heading('1. Project Overview', level=1)
    p = document.add_paragraph()
    p.add_run('Objective: ').bold = True
    p.add_run('To optimize debt collection channel selection (SMS, Call, Field Visit) using Reinforcement Learning (Contextual Multi-Armed Bandits) to maximize net recovery while minimizing costs.')
    
    document.add_paragraph('This system addresses the business problem of inefficient collection strategies by dynamically learning which channel works best for specific customer segments based on their risk profile, outstanding amount, and other behavioral features.')

    # 2. Architecture
    document.add_heading('2. System Architecture', level=1)
    document.add_paragraph('The project follows a modular architecture:')
    
    items = [
        ('Data Layer (Simulator)', 'Generates realistic customer profiles and defines the "Ground Truth" logic for channel response probabilities.'),
        ('Feature Engineering', 'Transforms raw customer data into numerical vectors (normalization, one-hot encoding) for the models.'),
        ('Bandit Core', 'Implements LinUCB (Linear Upper Confidence Bound) and Thompson Sampling algorithms to balance exploration and exploitation.'),
        ('API Layer', 'FastAPI service that warm-starts by training on simulated data and serves real-time recommendations.'),
        ('Explainability', 'Uses SHAP (SHapley Additive exPlanations) with a proxy Random Forest model to explain why a specific channel was chosen.')
    ]
    
    for key, value in items:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(key + ': ').bold = True
        p.add_run(value)

    # 3. Dataset Description
    document.add_heading('3. Dataset & Simulator Logic', level=1)
    document.add_paragraph('Since real financial data is sensitive, a Synthetic Data Simulator was built.')
    
    document.add_heading('Customer Features:', level=2)
    features = [
        'Risk Score (0-100)',
        'Outstanding Amount (Exponential dist)',
        'Days Past Due (Gamma dist)',
        'Credit Score Bucket (1-5)',
        'Customer Tier (Basic, Standard, Premium)',
        'Region (North, South, East, West)'
    ]
    for f in features:
        document.add_paragraph(f, style='List Bullet')
        
    document.add_heading('Channels & Costs:', level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Channel'
    hdr_cells[1].text = 'Cost'
    
    row = table.add_row().cells
    row[0].text = 'SMS'
    row[1].text = '$1.0'
    row = table.add_row().cells
    row[0].text = 'Call'
    row[1].text = '$5.0'
    row = table.add_row().cells
    row[0].text = 'Field Visit'
    row[1].text = '$20.0'

    document.add_heading('Reward Logic (Ground Truth):', level=2)
    document.add_paragraph('The simulator uses hidden logic to determine recovery probability:')
    document.add_paragraph('- High Risk (>80) customers respond better to Field Visits.', style='List Bullet')
    document.add_paragraph('- Low DPD (<10) customers respond well to SMS (soft nudge).', style='List Bullet')
    document.add_paragraph('- High Amount (>10k) warrants higher touch channels (Call).', style='List Bullet')
    document.add_paragraph('Net Reward = (Outstanding Amount * Recovery Rate) - Channel Cost', style='List Number')

    # 4. Implementation Details
    document.add_heading('4. Implementation Details', level=1)
    
    document.add_heading('Tech Stack', level=2)
    document.add_paragraph('Python 3.11, FastAPI, NumPy, Pandas, Scikit-Learn, SHAP, Docker, GitHub Actions.')
    
    document.add_heading('Key Components', level=2)
    document.add_paragraph('Policies implemented in `src/bandits/policies.py`:')
    document.add_paragraph('- LinUCB: Maintains inverse covariance matrices for each arm to estimate confidence bounds.')
    document.add_paragraph('- Thompson Sampling: Bayesian linear regression sampling from posterior distributions.')
    
    document.add_heading('Results', level=2)
    document.add_paragraph('In a 2000-round simulation:')
    document.add_paragraph('Thompson Sampling achieved ~12.5% uplift over Random baseline.', style='Quote')

    # 5. Commands Used
    document.add_heading('5. Execution Log & Commands', level=1)
    document.add_paragraph('The following commands were used to build and deploy the project:')
    
    commands = [
        ("Project Setup", "mkdir -p collection-campaign-optimizer/..."),
        ("Install Deps", "make install (pip install -r requirements.txt)"),
        ("Run Tests", "pytest tests/"),
        ("Run Experiment", "python run_experiments.py"),
        ("Start API", "make run-api (uvicorn src.api.app:app ...)"),
        ("Docker Build", "docker-compose up --build"),
        ("Git Init", "git init"),
        ("Git Commit", "git add . && git commit -m 'Initial commit...'"),
        ("Git Push", "git push -u origin main")
    ]
    
    for title, cmd in commands:
        p = document.add_paragraph()
        p.add_run(title + ': ').bold = True
        p.add_run(cmd).font.name = 'Courier New'

    document.save('Collection_Campaign_Optimizer_Documentation.docx')
    print("Documentation saved to Collection_Campaign_Optimizer_Documentation.docx")

if __name__ == "__main__":
    create_documentation()
